"""
Word document parser.
"""

from docx import Document as DocxDocument
from pathlib import Path

from core.ingestion.parser_base import BaseParser
from core.models.document import Document, DocumentMetadata, FileType, DocumentType


class WordParser(BaseParser):
    """Parser for Word documents (.docx)."""
    
    @property
    def supported_extensions(self) -> list[str]:
        return [".docx", ".doc"]
    
    @property
    def file_type(self) -> FileType:
        return FileType.WORD
    
    def parse(self, file_path: str) -> Document:
        """Parse a Word document into a Document."""
        doc = DocxDocument(file_path)
        
        paragraphs = []
        sections = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)
                # Detect headings as sections
                if para.style.name.startswith('Heading'):
                    sections.append(text)
        
        content = "\n\n".join(paragraphs)
        
        # Build metadata
        metadata = self._create_base_metadata(file_path)
        metadata.document_type = DocumentType.DOCUMENT
        metadata.sections = sections if sections else None
        
        return Document(
            id=self._generate_document_id(file_path),
            content=content,
            metadata=metadata,
        )
